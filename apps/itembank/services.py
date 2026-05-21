"""Service functions for item draft creation, versioning, and submission."""

import base64
import io
import json as _json
import logging
import shutil
import subprocess
import tempfile
import uuid
from datetime import timedelta

from django.contrib.auth import get_user_model
from django.core.exceptions import ObjectDoesNotExist, ValidationError
from django.core.files import File
from django.db import models, transaction
from django.utils import timezone
from django.conf import settings
from shared.storage import get_storage_backend
from decimal import Decimal, ROUND_HALF_UP
from collections import defaultdict

from .tasks import dispatch_item_status_notification

from .models import (
    Item,
    ItemTransition,
    ItemVersion,
    ItemComment,
    PanelVote,
    VaultAccess,
    VaultExportRequest,
    Paper,
    ItemUsage,
)
from apps.audit.models import AuditEvent
from workflow.guards import has_mandatory_metadata

logger = logging.getLogger(__name__)


@transaction.atomic
def create_or_update_item_draft(
    data: dict, author_auth: dict, item_id: str = None
) -> Item:
    """
    Create a new draft item or auto-save an existing item by creating a new version.

    Args:
        data: Item metadata and version content payload.
        author_auth: Auth payload containing the author's subject identifier.
        item_id: Existing item identifier for auto-save updates.

    Returns:
        The saved Item instance.
    """
    # Extract version-specific data that belongs on ItemVersion, not Item.
    content = data.pop("content", None)
    # Use None as the sentinel so omitted asset_refs do not wipe existing refs.
    asset_refs = data.pop("asset_refs", None)

    # Resolve author to a local User model instance (maps Keycloak `sub` to user.pk).
    User = get_user_model()
    try:
        author_user = User.objects.get(keycloak_sub=author_auth["sub"])
    except ObjectDoesNotExist as exc:
        raise ObjectDoesNotExist("Author user not found for provided auth sub") from exc

    if not item_id:
        # Create a brand new draft item.
        item = Item.objects.create(
            author_id=author_user,
            status=Item.Status.DRAFT,
            **data,
        )
        version_no = 1

        # Audit the initial draft creation so the item lifecycle starts with a
        # tamper-evident event in the platform audit log.
        AuditEvent.record(
            actor_id=author_auth["sub"],
            action="ITEM_DRAFT_CREATED",
            entity_type="item",
            entity_id=str(item.id),
            old_state=None,
            new_state={"status": Item.Status.DRAFT},
        )
    else:
        # Lock the existing item before updating it.
        item = Item.objects.select_for_update().get(id=item_id, author_id=author_user)

        if item.status not in [Item.Status.DRAFT, Item.Status.REVISED]:
            raise ValueError("You can only auto-save items in Draft or Revised states.")

        # Determine the next version number.
        last_version = item.versions.order_by("-version_no").first()
        version_no = last_version.version_no + 1 if last_version else 1

        if content is None:
            content = last_version.content if last_version else ""

        # If asset_refs omitted in an autosave, preserve the previous version refs.
        if asset_refs is None:
            asset_refs = last_version.asset_refs if last_version else []

        # Update the item-level metadata.
        for key, value in data.items():
            setattr(item, key, value)
        item.save()

    # Capture a metadata snapshot for the version record.
    metadata_snapshot = {
        "subject": item.subject,
        "topic": item.topic,
        "cognitive_level": item.cognitive_level,
        "difficulty": item.difficulty,
        "marks": str(item.marks) if item.marks is not None else None,
        "time": item.time,
        "source": item.source,
        "blueprint_ref": item.blueprint_ref,
    }

    # For new items, default asset_refs to empty list if none provided.
    if asset_refs is None:
        asset_refs = []

    new_version = ItemVersion.objects.create(
        item_id=item,
        version_no=version_no,
        content=content,
        metadata_snapshot=metadata_snapshot,
        asset_refs=asset_refs,
        saved_by=author_user,
    )

    # Link the item to the newly created active version.
    item.current_version_id = new_version.id
    item.save(update_fields=["current_version_id"])

    return item


@transaction.atomic
def submit_item_for_review(item_id: str, author_auth: dict) -> Item:
    """
    Submit a draft item for peer review after validating state and metadata.

    Args:
        item_id: The item identifier to submit.
        author_auth: Auth payload containing the author's subject identifier.

    Returns:
        The submitted Item instance.
    """
    # Resolve author to a local User model instance (maps Keycloak `sub` to user.pk).
    User = get_user_model()
    try:
        author_user = User.objects.get(keycloak_sub=author_auth["sub"])
    except ObjectDoesNotExist as exc:
        raise ObjectDoesNotExist("Author user not found for provided auth sub") from exc

    # Fetch and lock the item row for a safe state transition.
    item = Item.objects.select_for_update().get(id=item_id, author_id=author_user)

    # Ensure only draft-like states can be submitted.
    if item.status not in [Item.Status.DRAFT, Item.Status.REVISED]:
        raise ValidationError(
            f"Cannot submit an item currently in state: {item.status}"
        )

    # Require all mandatory metadata before submission.
    if not has_mandatory_metadata(item):
        raise ValidationError(
            "All mandatory metadata fields (subject, topic, cognitive level, difficulty, "
            "time, marks, source, syllabus reference) must be completed before submission."
        )

    old_status = item.status

    # Transition the item into the submitted state.
    item.status = Item.Status.SUBMITTED
    item.save(update_fields=["status"])

    # Record the workflow transition for history tracking.
    ItemTransition.objects.create(
        item_id=item,
        from_state=old_status,
        to_state=Item.Status.SUBMITTED,
        actor_id=author_user,
        justification="Item submitted for peer review",
    )

    # Record the system audit event.
    AuditEvent.record(
        actor_id=author_auth["sub"],
        action="ITEM_SUBMITTED",
        entity_type="item",
        entity_id=str(item.id),
        old_state={"status": old_status},
        new_state={"status": Item.Status.SUBMITTED},
    )

    return item


def scan_for_viruses(blob: bytes) -> bool:
    """Run a best-effort ClamAV scan against an uploaded blob."""

    clamscan = shutil.which("clamscan")
    if not clamscan:
        raise RuntimeError("Virus scanner unavailable: clamscan binary was not found.")

    with tempfile.NamedTemporaryFile(suffix=".upload", delete=True) as temp_file:
        temp_file.write(blob)
        temp_file.flush()
        try:
            completed = subprocess.run(
                [clamscan, "--no-summary", temp_file.name],
                capture_output=True,
                text=True,
                check=False,
                timeout=10,
            )
        except subprocess.TimeoutExpired as exc:
            logger.error("Virus scan timeout for temp file %s: %s", temp_file.name, exc)
            # Treat scan timeout as a failed (non-clean) result to be conservative
            return False

    if completed.returncode == 0:
        return True
    if completed.returncode == 1:
        return False

    raise RuntimeError(
        completed.stderr.strip() or completed.stdout.strip() or "Virus scan failed."
    )


def upload_to_vault_bucket(asset_ref: str, file_obj) -> str:
    """Persist an uploaded asset to the configured storage backend."""

    storage_backend = get_storage_backend()
    return storage_backend.save(asset_ref, File(file_obj))


def process_asset_upload(file_obj) -> str:
    """Process an uploaded asset file: virus-scan, store in vault, and return asset_ref.

    This is a lightweight, documented implementation used in development. In
    production the ClamAV scanner and object-storage helpers should be used by
    importing shared.security.scan_for_viruses and
    shared.storage.upload_to_vault_bucket.

    Steps performed:
    1. Minimal probe of the incoming file-like object to ensure it is readable
       (this also prevents the function from appearing unused by linters).
    2. (Placeholder) Run virus-scan logic. Currently mocked to pass.
    3. Generate a stable unique asset reference string.
    4. (Placeholder) Upload the file object to the blob vault.

    Args:
        file_obj: A file-like object (must support read() and seek()).

    Returns:
        A unique asset reference string to be stored on the ItemVersion.
    """

    # Probe the file to ensure the caller provided a readable object, then scan
    # and persist the upload before returning a vault reference.
    try:
        file_bytes = file_obj.read()
    except Exception as exc:
        raise ValueError(
            "Provided file_obj is not a readable file-like object"
        ) from exc

    # Ensure we have a fresh, seekable stream. If the original file-like
    # object does not support seek, recreate a BytesIO from the raw bytes.
    try:
        file_obj.seek(0)
    except (AttributeError, OSError, ValueError):
        file_obj = io.BytesIO(file_bytes)

    is_clean = scan_for_viruses(file_bytes)

    # Rewind or recreate the stream before upload to ensure the upload sees
    # the full content at position 0.
    try:
        file_obj.seek(0)
    except (AttributeError, OSError, ValueError):
        file_obj = io.BytesIO(file_bytes)

    if not is_clean:
        raise ValueError("File failed virus scan. Upload rejected and quarantined.")

    asset_ref = f"asset_{uuid.uuid4().hex}"

    try:
        upload_to_vault_bucket(asset_ref, file_obj)
    except Exception as exc:
        raise RuntimeError(
            f"Failed to upload asset {asset_ref} to vault storage."
        ) from exc

    return asset_ref


@transaction.atomic
def restore_item_version(item_id: str, version_id: str, actor_auth: dict) -> Item:
    """Restore a previous item version by creating a new version copy.

    This is a non-destructive restore operation: the historical
    ItemVersion identified by ``version_id`` is copied into a brand new
    ItemVersion record. The item's metadata fields are reverted to the
    values recorded in the historical version's ``metadata_snapshot`` and
    the item's ``current_version_id`` is updated to point to the newly
    created version.

    Preconditions / Validation:
    - The caller (``actor_auth["sub"]``) must be the item's assigned
      author.
    - The item must be in a state that allows restores ("Draft" or
      "Revised").

    Side effects:
    - Creates a new ItemVersion instance.
    - Updates and saves the Item instance.
    - Emits an AuditEvent recording the restore operation.

    Args:
        item_id: UUID/PK of the Item to act on.
        version_id: UUID/PK of the historical ItemVersion to copy.
        actor_auth: Auth information for the acting user; expects a
            "sub" key containing the user id.

    Returns:
        The updated Item instance (with current_version_id set to the
        newly created version).

    Raises:
        ValueError: If validation fails or the requested historical
            version does not exist.
    """
    User = get_user_model()
    try:
        resolved_user = User.objects.get(keycloak_sub=actor_auth["sub"])
    except ObjectDoesNotExist as exc:
        raise ValueError("Author user not found for provided auth sub") from exc

    try:
        item = Item.objects.select_for_update().get(id=item_id)
    except ObjectDoesNotExist as exc:
        raise ValueError("Item not found.") from exc

    # Validation Constraints
    if item.author_id_id != resolved_user.id:
        raise ValueError("Only the assigned author can restore item versions.")
    if item.status not in [Item.Status.DRAFT, Item.Status.REVISED]:
        raise ValueError(
            f"Cannot restore versions while item is in {item.status} state."
        )

    # Fetch historical snapshot
    try:
        historical_version = item.versions.get(id=version_id)
    except ObjectDoesNotExist as exc:
        raise ValueError("The requested version does not exist for this item.") from exc

    last_version = item.versions.order_by("-version_no").first()
    new_version_no = last_version.version_no + 1 if last_version else 1

    # Create the new version by perfectly copying the historical one
    new_version = ItemVersion.objects.create(
        item_id=item,
        version_no=new_version_no,
        content=historical_version.content,
        metadata_snapshot=historical_version.metadata_snapshot,
        asset_refs=historical_version.asset_refs,
        saved_by=resolved_user,
    )

    # Revert Item metadata to match the restored snapshot
    snapshot = historical_version.metadata_snapshot
    item.current_version_id = new_version.id
    item.subject = snapshot.get("subject", item.subject)
    item.topic = snapshot.get("topic", item.topic)
    item.cognitive_level = snapshot.get("cognitive_level", item.cognitive_level)
    item.difficulty = snapshot.get("difficulty", item.difficulty)
    item.time = snapshot.get("time", item.time)
    item.source = snapshot.get("source", item.source)
    item.blueprint_ref = snapshot.get("blueprint_ref", item.blueprint_ref)

    marks_val = snapshot.get("marks")
    item.marks = float(marks_val) if marks_val else None

    item.save()

    AuditEvent.record(
        actor_id=actor_auth["sub"],
        action="ITEM_VERSION_RESTORED",
        entity_type="item",
        entity_id=str(item.id),
        new_state={
            "restored_to_version": historical_version.version_no,
            "new_version_no": new_version_no,
        },
    )

    return item


@transaction.atomic
def process_suggestion_decision(
    item_id: str, suggestion_id: str, data: dict, actor_auth: dict
) -> dict:
    """Process a decision on an inline suggestion (accept or decline).

    The system stores inline suggestions as ItemComment records. This
    function resolves a suggestion by marking it "resolved" and
    optionally creates a rationale reply record also stored as an
    ItemComment (linked via the ``anchor_path``).

    Validation rules:
    - Only the Item's assigned author may accept or decline suggestions.
    - Suggestions may only be processed while the item is in
      "In Review" or "Revised" states.
    - Suggestions that are already resolved cannot be processed again.

    Args:
        item_id: UUID/PK of the Item being acted on.
        suggestion_id: UUID/PK of the ItemComment representing the
            suggestion.
        data: Dictionary containing the decision payload. Expected keys
            include "decision" (e.g. "accept" or "decline") and
            optionally "rationale" (a free-text explanation).
        actor_auth: Auth information for the acting user; expects a
            "sub" key containing the user id.

    Returns:
        A dict summarising the outcome containing the suggestion id,
        its updated status, and the id of any rationale reply that was
        created.

    Raises:
        ValueError: If validation fails or the suggestion cannot be
            located.
    """
    User = get_user_model()
    try:
        resolved_user = User.objects.get(keycloak_sub=actor_auth["sub"])
    except ObjectDoesNotExist as exc:
        raise ValueError("Author user not found for provided auth sub") from exc

    try:
        item = Item.objects.select_for_update().get(id=item_id)
    except ObjectDoesNotExist as exc:
        raise ValueError("Item not found.") from exc

    try:
        suggestion = ItemComment.objects.get(
            id=suggestion_id, item_version_id__item_id=item
        )
    except ObjectDoesNotExist as exc:
        raise ValueError("Suggestion not found.") from exc

    # RBAC/State Validation
    if item.author_id_id != resolved_user.id:
        raise ValueError("Only the Item Writer can accept or decline suggestions.")
    if item.status not in [Item.Status.IN_REVIEW, Item.Status.REVISED]:
        raise ValueError(
            f"Cannot process suggestions while item is in {item.status} state."
        )
    if suggestion.status == "resolved":
        raise ValueError("This suggestion has already been resolved.")

    suggestion.status = "resolved"
    suggestion.save(update_fields=["status"])

    # using the ItemComment table, leveraging the anchor_path to link it(If declined (or if a rationale was provided for an accept))
    rationale_record = None
    if data.get("rationale"):
        rationale_record = ItemComment.objects.create(
            item_version_id=suggestion.item_version_id,
            anchor_path=f"reply_to_{suggestion.id}",  # Links the rationale to the original suggestion
            body=f"[{data['decision'].upper()}] Rationale: {data['rationale']}",
            status="resolved",  # Replies are born resolved so they don't clutter the open queue
            created_by=resolved_user,
        )

    AuditEvent.record(
        actor_id=actor_auth["sub"],
        action=f"SUGGESTION_{data['decision'].upper()}",
        entity_type="item_comment",
        entity_id=str(suggestion.id),
        new_state={"rationale_provided": bool(data.get("rationale"))},
    )

    return {
        "suggestion_id": str(suggestion.id),
        "status": suggestion.status,
        "rationale_id": str(rationale_record.id) if rationale_record else None,
    }


def check_and_escalate_overdue_reviews() -> int:
    """Find and escalate items that have exceeded the review SLA.

    This routine computes a cutoff datetime representing 5 business days
    prior to now (weekdays only), then scans items currently in the
    "In Review" status. For each item it determines when the item
    entered the "In Review" state by inspecting recorded transitions;
    if that timestamp is older than the cutoff the item is considered
    to have breached the SLA. Each breach is logged and an AuditEvent
    is recorded. The function returns the total number of escalated
    items.

    Returns:
        int: Number of items escalated due to SLA breach.
    """
    # 5 business days by counting weekdays backward.
    sla_cutoff = timezone.now()
    business_days = 0

    while business_days < 5:
        sla_cutoff -= timedelta(days=1)
        if sla_cutoff.weekday() < 5:
            business_days += 1

    # Only look at items currently stuck in the review queue
    stuck_items = Item.objects.filter(status=Item.Status.IN_REVIEW)
    escalated_count = 0

    for item in stuck_items:
        # Find the exact moment this item entered the 'In Review' state
        entry_transition = (
            item.transitions.filter(to_state=Item.Status.IN_REVIEW)
            .order_by("-occurred_at")
            .first()
        )

        if entry_transition and entry_transition.occurred_at < sla_cutoff:
            # The SLA is breached.
            # 1. Log the breach for the system administrators
            logger.warning(
                "SLA BREACH: Item %s has been In Review since %s.",
                item.id,
                entry_transition.occurred_at,
            )

            # Avoid duplicate SLA escalation audit events within the recent window
            recent_since = timezone.now() - timedelta(days=1)
            already_recorded = AuditEvent.objects.filter(
                actor_id="SYSTEM",
                action="SLA_ESCALATION_TRIGGERED",
                entity_type="item",
                entity_id=str(item.id),
                created_at__gte=recent_since,
            ).exists()

            if not already_recorded:
                AuditEvent.record(
                    actor_id="SYSTEM",  # Triggered by the server, not a user
                    action="SLA_ESCALATION_TRIGGERED",
                    entity_type="item",
                    entity_id=str(item.id),
                    new_state={
                        "days_overdue": (
                            timezone.now() - entry_transition.occurred_at
                        ).days
                    },
                )

                escalated_count += 1

    return escalated_count


@transaction.atomic
def register_panel_vote(
    item_id: str, panellist_id: str, vote_type: str, justification: str
) -> Item:
    """
    Casts an item panel vote, checks word thresholds,
    and evaluates consensus state changes automatically.
    """
    item = Item.objects.select_for_update().get(id=item_id)

    if item.status != Item.Status.MODERATION_PANEL:
        raise ValueError(
            f"Item is not in the moderation phase. Current status: {item.status}"
        )

    # Enforce word limit constraint
    word_count = len(justification.split())
    if word_count < 30:
        raise ValueError(
            f"Justification narrative must be at least 30 words. Current count: {word_count}"
        )

    # Prevent duplicate panellist votes (unique_together enforced at DB level).
    if PanelVote.objects.filter(item_id=item, panellist_id=panellist_id).exists():
        raise ValueError("This panellist has already voted on this item.")

    # Save the vote
    PanelVote.objects.create(
        item_id=item,
        panellist_id=panellist_id,
        vote=vote_type,
        justification=justification,
    )

    # Evaluate current matching votes
    votes = item.panel_votes.all()
    approve_count = votes.filter(vote="Approve").count()
    reject_count = votes.filter(vote="Reject").count()

    # Consensus rules (Default 2 out of 3 match wins)
    if approve_count >= 2:
        # Record the SRS-defined approval step before auto-locking the item for use.
        previous_status = item.status
        item.status = Item.Status.APPROVED
        item.save(update_fields=["status"])

        ItemTransition.objects.create(
            item_id=item,
            from_state=previous_status,
            to_state=Item.Status.APPROVED,
            actor_id=panellist_id,
            justification=justification,
        )

        item.status = Item.Status.LOCKED_FOR_USE
        item.save(update_fields=["status"])

        ItemTransition.objects.create(
            item_id=item,
            from_state=Item.Status.APPROVED,
            to_state=Item.Status.LOCKED_FOR_USE,
            actor_id=panellist_id,
            justification="Item approved and automatically locked for use",
        )

        # Fire async notification only after the transaction commits.
        transaction.on_commit(
            lambda item_id=str(item.id), author_id=str(item.author_id_id): (
                dispatch_item_status_notification.delay(
                    item_id,
                    author_id,
                    Item.Status.LOCKED_FOR_USE,
                )
            )
        )

        # Log forensic snapshot to security audit
        AuditEvent.record(
            actor_id="SYSTEM",
            action="ITEM_CONSENSUAL_APPROVAL",
            entity_type="item",
            entity_id=str(item.id),
            new_state={
                "status": Item.Status.LOCKED_FOR_USE,
                "approvers": [
                    str(v.panellist_id) for v in votes.filter(vote="Approve")
                ],
            },
        )
    elif reject_count >= 2:
        previous_status = item.status
        item.status = Item.Status.REJECTED
        item.save(update_fields=["status"])

        # Record workflow transition for the consensual rejection
        ItemTransition.objects.create(
            item_id=item,
            from_state=previous_status,
            to_state=Item.Status.REJECTED,
            actor_id=panellist_id,
            justification="Consensual panel rejection",
        )

        # Consolidate rejection rationale for the 5-minute notification trigger
        consolidated_rationale = [v.justification for v in votes.filter(vote="Reject")]
        transaction.on_commit(
            lambda item_id=str(item.id), author_id=str(item.author_id_id), rationales=consolidated_rationale: (
                dispatch_item_status_notification.delay(
                    item_id,
                    author_id,
                    Item.Status.REJECTED,
                    rationales=rationales,
                )
            )
        )

        AuditEvent.record(
            actor_id="SYSTEM",
            action="ITEM_CONSENSUAL_REJECTION",
            entity_type="item",
            entity_id=str(item.id),
            new_state={"status": Item.Status.REJECTED},
        )

    return item


@transaction.atomic
def execute_vault_cosign(request_id: str, cosigner_id: str) -> VaultExportRequest:
    """
    Co-signs and executes an export request with strict anti-circumvention validations.
    """
    User = get_user_model()
    try:
        cosigner_user = User.objects.get(keycloak_sub=cosigner_id)
    except ObjectDoesNotExist as exc:
        raise ValueError("Cosigner user not found for provided auth sub.") from exc

    req = VaultExportRequest.objects.select_for_update().get(id=request_id)

    if req.status != "Pending":
        raise ValueError(
            f"This export request is no longer active. Status: {req.status}"
        )
    if timezone.now() > req.expires_at:
        req.status = "Expired"
        req.save()
        raise ValueError("The 72-hour validation window for this request has expired.")

    # Prevent self-signing anti-circumvention rule
    if req.requester_id_id == cosigner_user.id:
        raise ValueError(
            "Security Boundary Exception: Initiating officer cannot act as the co-signing authoriser."
        )

    # Execute request
    req.cosigner_id = cosigner_user
    req.status = "Executed"
    req.save()

    # Record the actual physical vault access log line
    AuditEvent.record(
        actor_id=str(cosigner_id),
        action="VAULT_CONTENT_EXPORTED",
        entity_type="vault",
        entity_id=str(req.id),
        new_state={"scope_length": len(req.scope)},
    )

    return req


def _marks_tolerance() -> Decimal:
    """Configured per-paper marks tolerance (Decimal)."""
    return Decimal(str(getattr(settings, "PAPER_MARKS_TOLERANCE", "0")))


def _recent_sittings(limit: int, exclude: str | None = None) -> list[str]:
    """Return the most-recent N distinct sitting refs from ``ItemUsage``.

    Uses ``Max(recorded_at)`` per sitting so the query is portable across
    PostgreSQL and SQLite — the previous ``values_list().distinct()
    .order_by("-recorded_at")`` form fails on PostgreSQL because the
    ORDER BY column is not in the SELECT DISTINCT list.
    """
    qs = ItemUsage.objects.values("sitting_ref")
    if exclude:
        qs = qs.exclude(sitting_ref=exclude)
    qs = qs.annotate(latest=models.Max("recorded_at")).order_by("-latest")
    return list(qs.values_list("sitting_ref", flat=True)[:limit])


def get_cooled_down_items(sitting_ref: str | None = None) -> set:
    """Return the set of item IDs in the cool-down window."""
    cool_down = getattr(settings, "ITEM_COOLDOWN_SITTINGS", 3)
    recent = _recent_sittings(cool_down, exclude=sitting_ref)
    if not recent:
        return set()
    return set(
        ItemUsage.objects.filter(sitting_ref__in=recent).values_list(
            "item_id", flat=True
        )
    )


def _validate_blueprint(items, *, blueprint_ref: str, total_marks: Decimal) -> None:
    """Enforce topic-coverage rules declared by ``NBES_BLUEPRINTS[blueprint_ref]``.

    Raises ``ValueError`` with the SRS-mandated error message when a topic
    is under-represented compared with the blueprint percentage.
    """
    if not blueprint_ref:
        return
    catalogue = getattr(settings, "NBES_BLUEPRINTS", {}) or {}
    blueprint = catalogue.get(blueprint_ref)
    if not blueprint:
        return  # Unknown blueprint refs are tolerated until Phase 4 lands.

    topic_marks: dict[str, Decimal] = defaultdict(lambda: Decimal("0"))
    for item in items:
        if item.topic and item.marks is not None:
            topic_marks[item.topic] += Decimal(str(item.marks))

    required = blueprint.get("topics") or {}
    tolerance = _marks_tolerance()
    for topic, pct in required.items():
        required_marks = (total_marks * Decimal(str(pct)) / Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        actual = topic_marks.get(topic, Decimal("0"))
        if actual + tolerance < required_marks:
            raise ValueError(
                f"Topic coverage does not satisfy the blueprint constraints: "
                f"Topic {topic} under-represented vs blueprint "
                f"(required {required_marks}, got {actual})."
            )


def _validate_sections(items, *, sections: list, blueprint_ref: str) -> None:
    """Enforce the SRS section-structure rule.

    A submitted ``sections`` payload must:
      * cover every item in the paper exactly once
      * each section's ``marks`` (when supplied) must match Σ item.marks
      * align with the blueprint's section catalogue when defined
    """
    if not sections:
        # If the blueprint declares sections, the payload must too.
        catalogue = getattr(settings, "NBES_BLUEPRINTS", {}) or {}
        if catalogue.get(blueprint_ref, {}).get("sections"):
            raise ValueError(
                "Blueprint Violation: blueprint requires section structure; "
                "payload supplied no sections."
            )
        return

    items_by_id = {str(item.id): item for item in items}
    seen: set[str] = set()
    for idx, section in enumerate(sections, start=1):
        section_item_ids = [str(i) for i in section.get("item_ids", [])]
        if not section_item_ids:
            raise ValueError(f"Section {idx} contains no items.")
        for item_id in section_item_ids:
            if item_id in seen:
                raise ValueError(f"Item {item_id} appears in more than one section.")
            if item_id not in items_by_id:
                raise ValueError(
                    f"Section {idx} references item {item_id} which is not in the paper."
                )
            seen.add(item_id)
        declared_marks = section.get("marks")
        if declared_marks is not None:
            actual = sum(
                Decimal(str(items_by_id[i].marks or 0)) for i in section_item_ids
            )
            if Decimal(str(declared_marks)) != actual:
                raise ValueError(
                    f"Section {idx} marks mismatch: declared "
                    f"{declared_marks}, computed {actual}."
                )

    if seen != set(items_by_id):
        missing = set(items_by_id) - seen
        raise ValueError(
            f"Section structure does not cover every item; missing {sorted(missing)}."
        )

    catalogue = getattr(settings, "NBES_BLUEPRINTS", {}) or {}
    blueprint_sections = (catalogue.get(blueprint_ref) or {}).get("sections") or []
    if blueprint_sections and len(blueprint_sections) != len(sections):
        raise ValueError(
            f"Blueprint Violation: blueprint defines {len(blueprint_sections)} "
            f"sections; payload supplied {len(sections)}."
        )


def _validate_time_allocation(items, *, time_limit: int) -> None:
    """Σ item.time must be ≤ paper.time_limit (when both are populated)."""
    total = sum((item.time or 0) for item in items)
    if total > int(time_limit):
        raise ValueError(
            f"Time allocation exceeds paper time limit: items sum to {total} "
            f"but paper limit is {time_limit}."
        )


@transaction.atomic
def create_manual_paper(data: dict, user, request=None) -> Paper:
    """Construct a paper manually from a curated list of locked items.

    Implements SRS-NBE-F02-08 manual-construction rules: marks total,
    blueprint coverage, section structure, time allocation, cool-down,
    and per-item vault-read logging.
    """
    item_ids = list(data["item_ids"])
    sitting_ref = data["sitting_ref"]
    subject = data["subject"]
    mode = data["mode"]
    total_marks = Decimal(str(data["total_marks"]))
    time_limit = int(data["time_limit"])
    blueprint_ref = data.get("blueprint_ref") or ""
    sections = data.get("sections") or []

    if len(set(str(i) for i in item_ids)) != len(item_ids):
        raise ValueError("Duplicate item_ids in payload.")

    items_qs = Item.objects.filter(id__in=item_ids, status=Item.Status.LOCKED_FOR_USE)
    items_by_id = {str(item.id): item for item in items_qs}
    if len(items_by_id) != len(item_ids):
        missing = [str(i) for i in item_ids if str(i) not in items_by_id]
        raise ValueError(
            "Paper construction rejects items in Draft, Submitted, or Rejected states: "
            f"items not available (must be Locked for Use): {missing}"
        )
    items = [items_by_id[str(i)] for i in item_ids]  # preserve caller order

    actual_marks = sum(
        (Decimal(str(item.marks)) for item in items if item.marks is not None),
        Decimal("0"),
    )
    if actual_marks != total_marks:
        raise ValueError(
            f"Total marks ({actual_marks}) do not match the configured paper total ({total_marks})."
        )

    if any(item.subject != subject for item in items):
        raise ValueError("All items must belong to the same subject as the paper.")

    _validate_time_allocation(items, time_limit=time_limit)
    _validate_blueprint(items, blueprint_ref=blueprint_ref, total_marks=total_marks)
    _validate_sections(items, sections=sections, blueprint_ref=blueprint_ref)

    cool_down = getattr(settings, "ITEM_COOLDOWN_SITTINGS", 3)
    recent = _recent_sittings(cool_down, exclude=sitting_ref)
    if recent:
        for item in items:
            if ItemUsage.objects.filter(item_id=item, sitting_ref__in=recent).exists():
                raise ValueError(
                    f"Item {item.id} has been used within the cool-down window."
                )

    paper = Paper.objects.create(
        sitting_ref=sitting_ref,
        subject=subject,
        mode=mode,
        total_marks=total_marks,
        time_limit=time_limit,
        item_ids=[str(item.id) for item in items],
        sections=sections,
        blueprint_ref=blueprint_ref,
        status=Paper.Status.CONSTRUCTED,
    )

    for item in items:
        ItemUsage.objects.create(item_id=item, sitting_ref=sitting_ref, count=1)

    log_vault_reads(items, user=user, kind="read", request=request)

    AuditEvent.record(
        actor_id=user.keycloak_sub,
        action="PAPER_CONSTRUCTED",
        entity_type="paper",
        entity_id=str(paper.id),
        new_state={
            "sitting_ref": sitting_ref,
            "item_count": len(items),
            "blueprint_ref": blueprint_ref,
            "has_sections": bool(sections),
        },
    )
    return paper


@transaction.atomic
def submit_paper_for_approval(paper_id: str, user) -> Paper:
    """Transition a constructed paper into NBEC approval (SRS-NBE-F02-08).

    Mirrors the F02-08 acceptance criterion that a complete paper can be
    "saved and submitted for NBEC approval".
    """
    paper = Paper.objects.select_for_update().get(id=paper_id)
    if paper.status not in (Paper.Status.CONSTRUCTED, Paper.Status.DRAFT):
        raise ValueError(f"Cannot submit paper in state {paper.status} for approval.")
    previous = paper.status
    paper.status = Paper.Status.READY_FOR_APPROVAL
    paper.save(update_fields=["status", "updated_at"])
    AuditEvent.record(
        actor_id=user.keycloak_sub,
        action="PAPER_SUBMITTED_FOR_APPROVAL",
        entity_type="paper",
        entity_id=str(paper.id),
        old_state={"status": previous},
        new_state={"status": paper.status},
    )
    return paper


def log_vault_reads(items, *, user, kind: str = "read", request=None) -> None:
    """Record ``VaultAccess`` rows + per-item ``VAULT_READ`` audit events.

    SRS-NBE-F02-07 requires "vault read for paper construction logged per
    item per user per session" with the audit entry reaching System 22
    within 1 minute. We:
      * write one ``VaultAccess`` row per item with session_id + IP
      * emit a per-item ``AuditEvent`` so the chain hash + outbox push the
        record into System 22 via the existing pipeline.
    """
    session_id = None
    ip = None
    if request is not None:
        session_id = getattr(request, "request_id", None)
        ip = getattr(request, "ip_address", None) or request.META.get("REMOTE_ADDR")
    action = "VAULT_EXPORTED" if kind == "export" else "VAULT_READ"
    for item in items:
        item_instance = item if isinstance(item, Item) else None
        if item_instance is None:
            try:
                item_instance = Item.objects.get(id=item)
            except ObjectDoesNotExist:
                continue
        VaultAccess.objects.create(
            item_id=item_instance,
            actor_id=user,
            kind=kind,
            session_id=str(session_id) if session_id else None,
            ip=str(ip) if ip else None,
        )
        AuditEvent.record(
            actor_id=user.keycloak_sub,
            action=action,
            entity_type="item",
            entity_id=item_instance.id,
            new_state={
                "kind": kind,
                "session_id": str(session_id) if session_id else None,
                "ip": str(ip) if ip else None,
            },
            ip_address=str(ip) if ip else None,
        )


def _select_items_for_constraints(
    pool: list,
    diff_dist: dict,
    topic_dist: dict,
    total_marks: Decimal,
    exclude_ids: set | None = None,
) -> tuple[list, Decimal, list]:
    """Greedy bucket selector returning (items, total_marks, errors)."""
    exclude_ids = exclude_ids or set()
    buckets: dict = defaultdict(list)
    for item in pool:
        if item.id in exclude_ids:
            continue
        if not item.difficulty or not item.topic or item.marks is None:
            continue
        if item.difficulty in diff_dist and item.topic in topic_dist:
            buckets[(item.difficulty, item.topic)].append(item)

    for bucket_items in buckets.values():
        bucket_items.sort(key=lambda x: x.marks or 0, reverse=True)

    required_marks_per_difficulty: dict = {
        diff: (total_marks * Decimal(str(diff_pct)) / Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        for diff, diff_pct in diff_dist.items()
    }
    required_marks_per_topic: dict = {
        topic: (total_marks * Decimal(str(topic_pct)) / Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        for topic, topic_pct in topic_dist.items()
    }

    selected: list = []
    selected_marks = Decimal("0")
    errors: list[str] = []

    while True:
        best_bucket = None
        best_score = Decimal("0")
        for (diff, topic), available in buckets.items():
            if not available:
                continue
            score = required_marks_per_difficulty.get(
                diff, Decimal("0")
            ) + required_marks_per_topic.get(topic, Decimal("0"))
            if score > best_score:
                best_score = score
                best_bucket = (diff, topic)

        if best_bucket is None or best_score <= 0:
            break

        diff, topic = best_bucket
        item = buckets[(diff, topic)].pop(0)
        item_marks = Decimal(str(item.marks))
        selected.append(item)
        selected_marks += item_marks
        required_marks_per_difficulty[diff] = max(
            Decimal("0"),
            required_marks_per_difficulty.get(diff, Decimal("0")) - item_marks,
        )
        required_marks_per_topic[topic] = max(
            Decimal("0"),
            required_marks_per_topic.get(topic, Decimal("0")) - item_marks,
        )

    target_marks_per_difficulty = {
        diff: (total_marks * Decimal(str(diff_pct)) / Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        for diff, diff_pct in diff_dist.items()
    }
    target_marks_per_topic = {
        topic: (total_marks * Decimal(str(topic_pct)) / Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        for topic, topic_pct in topic_dist.items()
    }

    for diff, remaining in required_marks_per_difficulty.items():
        if remaining > 0:
            errors.append(
                f"Difficulty {diff} remains under-represented vs blueprint "
                f"(required {target_marks_per_difficulty[diff]}, got {target_marks_per_difficulty[diff] - remaining})."
            )
    for topic, remaining in required_marks_per_topic.items():
        if remaining > 0:
            errors.append(
                f"Topic {topic} remains under-represented vs blueprint "
                f"(required {target_marks_per_topic[topic]}, got {target_marks_per_topic[topic] - remaining})."
            )
    return selected, selected_marks, errors


@transaction.atomic
def generate_paper_rule_based(data: dict, user, request=None) -> Paper:
    """Rule-based paper generation (SRS-NBE-F02-08).

    Honours difficulty distribution, topic coverage, marks allocation and
    cool-down. Produces ``variants_count`` variants when requested — each
    variant is a disjoint item set drawn from the same pool, matching
    SRS-NBE-F03-05.
    """
    sitting_ref = data["sitting_ref"]
    subject = data["subject"]
    mode = data["mode"]
    total_marks = Decimal(str(data["total_marks"]))
    time_limit = int(data["time_limit"])
    diff_dist = data["difficulty_distribution"]
    topic_dist = data["topic_coverage"]
    blueprint_ref = data.get("blueprint_ref", "") or ""
    variants_count = int(data.get("variants_count", 1) or 1)

    if abs(sum(diff_dist.values()) - 100) > 0.1:
        raise ValueError("Difficulty distribution percentages must sum to 100.")
    if abs(sum(topic_dist.values()) - 100) > 0.1:
        raise ValueError("Topic coverage percentages must sum to 100.")
    if variants_count < 1:
        raise ValueError("variants_count must be >= 1.")

    cooled_ids = get_cooled_down_items(sitting_ref)
    pool = list(
        Item.objects.filter(subject=subject, status=Item.Status.LOCKED_FOR_USE).exclude(
            id__in=cooled_ids
        )
    )

    primary, primary_marks, primary_errors = _select_items_for_constraints(
        pool, diff_dist, topic_dist, total_marks
    )
    if primary_errors:
        raise ValueError(
            "Unable to meet all topic/difficulty constraints. "
            + " ".join(primary_errors)
        )

    _validate_blueprint(primary, blueprint_ref=blueprint_ref, total_marks=total_marks)
    _validate_time_allocation(primary, time_limit=time_limit)

    tolerance = _marks_tolerance()
    if abs(primary_marks - total_marks) > tolerance:
        raise ValueError(
            f"Selected items total {primary_marks} deviates from required "
            f"{total_marks} by more than allowed tolerance {tolerance}."
        )

    used_ids = {item.id for item in primary}
    variant_payloads: list[dict] = []
    for variant_no in range(2, variants_count + 1):
        variant_items, variant_marks, variant_errors = _select_items_for_constraints(
            pool, diff_dist, topic_dist, total_marks, exclude_ids=used_ids
        )
        if variant_errors:
            raise ValueError(
                f"Variant {variant_no}: unable to satisfy blueprint with "
                f"remaining pool. " + " ".join(variant_errors)
            )
        if abs(variant_marks - total_marks) > tolerance:
            raise ValueError(
                f"Variant {variant_no} marks {variant_marks} deviate from "
                f"target {total_marks} by more than tolerance {tolerance}."
            )
        _validate_blueprint(
            variant_items, blueprint_ref=blueprint_ref, total_marks=total_marks
        )
        _validate_time_allocation(variant_items, time_limit=time_limit)
        variant_payloads.append(
            {
                "variant_no": variant_no,
                "item_ids": [str(item.id) for item in variant_items],
                "total_marks": str(variant_marks),
            }
        )
        used_ids.update(item.id for item in variant_items)

    paper = Paper.objects.create(
        sitting_ref=sitting_ref,
        subject=subject,
        mode=mode,
        total_marks=total_marks,
        time_limit=time_limit,
        item_ids=[str(item.id) for item in primary],
        variants=variant_payloads,
        blueprint_ref=blueprint_ref,
        status=Paper.Status.CONSTRUCTED,
    )

    all_items = list(primary)
    pool_by_id = {str(item.id): item for item in pool}
    for variant in variant_payloads:
        for item_id in variant["item_ids"]:
            item = pool_by_id.get(item_id)
            if item is not None:
                all_items.append(item)

    for item in all_items:
        ItemUsage.objects.create(item_id=item, sitting_ref=sitting_ref, count=1)

    log_vault_reads(all_items, user=user, kind="read", request=request)

    AuditEvent.record(
        actor_id=user.keycloak_sub,
        action="PAPER_GENERATED",
        entity_type="paper",
        entity_id=str(paper.id),
        new_state={
            "sitting_ref": sitting_ref,
            "item_count": len(primary),
            "variants": len(variant_payloads),
            "total_marks": str(primary_marks),
        },
    )
    return paper


def _ordered_paper_items(paper: Paper) -> list:
    """Return the items referenced by ``paper.item_ids`` in their stored order."""
    by_id = {str(item.id): item for item in Item.objects.filter(id__in=paper.item_ids)}
    ordered = []
    for ref in paper.item_ids:
        item = by_id.get(str(ref))
        if item is not None:
            ordered.append(item)
    return ordered


def _item_body(item: Item) -> str:
    """Best-effort extraction of an item's renderable body text.

    The current_version's content may be plain text, HTML, or a JSON
    payload produced by the rich-text editor. Returns a single string
    safe to embed inside the export templates.
    """
    if not item.current_version_id:
        return ""
    try:
        version = item.versions.get(id=item.current_version_id)
    except ObjectDoesNotExist:
        return ""
    content = version.content or ""
    try:
        parsed = _json.loads(content)
    except (TypeError, ValueError):
        return content
    if isinstance(parsed, dict):
        # Common rich-text editor shape: {"prompt": "...", "options": [...]}.
        prompt = parsed.get("prompt") or parsed.get("body") or parsed.get("text") or ""
        options = parsed.get("options") or []
        if options:
            lines = [str(prompt)] if prompt else []
            for idx, opt in enumerate(options):
                if isinstance(opt, dict):
                    label = opt.get("label") or opt.get("text") or ""
                else:
                    label = str(opt)
                lines.append(f"  ({chr(ord('A') + idx)}) {label}")
            return "\n".join(lines)
        return str(prompt) if prompt else content
    return content


def _qr_code_png_base64(payload: str) -> str:
    """Return a base64-encoded PNG QR code for ``payload``.

    Imported lazily so a missing ``qrcode`` library never blocks Django
    boot — only callers actually generating exports pay the cost.
    """
    import qrcode  # local import

    img = qrcode.make(payload)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("ascii")


def _verification_url(paper: Paper) -> str:
    base = getattr(settings, "PAPER_VERIFICATION_BASE_URL", "https://verify.gsl.edu.gh")
    return f"{base.rstrip('/')}/paper/{paper.id}"


def _render_paper_payload(paper: Paper) -> tuple:
    """Return ``(items, item_records)`` where each record is a dict ready
    to feed any of the export formats."""
    items = _ordered_paper_items(paper)
    records = []
    for idx, item in enumerate(items, start=1):
        records.append(
            {
                "position": idx,
                "question_number": idx,
                "id": str(item.id),
                "item_type": item.item_type,
                "subject": item.subject,
                "topic": item.topic,
                "difficulty": item.difficulty,
                "cognitive_level": item.cognitive_level,
                "marks": str(item.marks) if item.marks is not None else None,
                "time": item.time,
                "body": _item_body(item),
            }
        )
    return items, records


def export_paper_pdf(paper: Paper) -> bytes:
    """Render the paper as a PDF using WeasyPrint.

    The cover sheet embeds a QR code pointing at the verification URL
    plus the explicit item count required by SRS-NBE-F02-08. If the
    paper carries a section structure the questions are grouped by
    section in the rendered PDF.
    """
    from django.template.loader import render_to_string
    from weasyprint import HTML  # local import — heavyweight dependency

    _, records = _render_paper_payload(paper)
    record_by_id = {r["id"]: r for r in records}

    sections_ctx: list = []
    for section in paper.sections or []:
        sections_ctx.append(
            {
                "name": section.get("name"),
                "marks": section.get("marks"),
                "time": section.get("time"),
                "items": [
                    record_by_id[str(item_id)]
                    for item_id in section.get("item_ids", [])
                    if str(item_id) in record_by_id
                ],
            }
        )

    qr_code = _qr_code_png_base64(_verification_url(paper))
    html_string = render_to_string(
        "itembank/paper.html",
        {
            "paper": paper,
            "items": records,
            "sections": sections_ctx,
            "qr_code": qr_code,
        },
    )
    return HTML(string=html_string).write_pdf()


def export_paper_docx(paper: Paper) -> bytes:
    """Render the paper as a Microsoft Word ``.docx`` document.

    Requires ``python-docx`` to be installed in the runtime environment.
    """
    try:
        from docx import Document  # local import
    except ImportError as exc:  # pragma: no cover - runtime guard
        raise RuntimeError(
            "python-docx must be installed to export papers in Word format."
        ) from exc

    _, records = _render_paper_payload(paper)
    document = Document()
    document.add_heading(f"{paper.subject}", level=0)
    document.add_paragraph(f"Sitting: {paper.sitting_ref}")
    document.add_paragraph(f"Mode: {paper.mode}")
    document.add_paragraph(f"Total marks: {paper.total_marks}")
    document.add_paragraph(f"Time limit: {paper.time_limit} minutes")
    if paper.blueprint_ref:
        document.add_paragraph(f"Blueprint: {paper.blueprint_ref}")
    document.add_paragraph(f"Paper ID: {paper.id}")
    document.add_paragraph(f"Verification: {_verification_url(paper)}")

    document.add_heading("Questions", level=1)
    for record in records:
        header = (
            f"Q{record['position']} — {record['marks'] or '—'} marks "
            f"({record['difficulty'] or '—'})"
        )
        document.add_heading(header, level=2)
        document.add_paragraph(record["body"] or "")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def export_paper_digital(paper: Paper) -> dict:
    """Return a JSON-serialisable digital export of the paper.

    System 10B (digital delivery) consumes this structure directly.
    """
    _, records = _render_paper_payload(paper)
    return {
        "paper_id": str(paper.id),
        "sitting_ref": paper.sitting_ref,
        "subject": paper.subject,
        "mode": paper.mode,
        "total_marks": str(paper.total_marks),
        "time_limit": paper.time_limit,
        "blueprint_ref": paper.blueprint_ref,
        "status": paper.status,
        "verification_url": _verification_url(paper),
        "items": records,
    }


def record_paper_export(paper: Paper, *, user, fmt: str, request=None) -> None:
    """Log a vault export (one row per item) plus a top-level audit event.

    Each format (PDF/Word/Digital) is treated as an export of the underlying
    items, so we log a ``VaultAccess`` row per item and a paper-level audit
    event for traceability.
    """
    items = _ordered_paper_items(paper)
    log_vault_reads(items, user=user, kind="export", request=request)
    AuditEvent.record(
        actor_id=user.keycloak_sub,
        action="PAPER_EXPORTED",
        entity_type="paper",
        entity_id=str(paper.id),
        new_state={"format": fmt, "item_count": len(items)},
    )
